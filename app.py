import streamlit as st
import pandas as pd
import anthropic
from docx import Document
from io import BytesIO

# --- Page Config ---
st.set_page_config(page_title="CRE IC Memo Generator", page_icon="🏢", layout="wide")
st.title("🏢 CRE IC Memo Generator (Smart Search)")
st.markdown("Upload a client's financial model. The app will auto-detect key tables (T-12, Unit Mix, etc.) and generate a formatted memo.")

# --- API Key Handling ---
try:
    api_key = st.secrets["ANTHROPIC_API_KEY"]
except (KeyError, FileNotFoundError):
    api_key = st.text_input("Enter your Anthropic API Key:", type="password")

# --- Inputs ---
col1, col2 = st.columns(2)
with col1:
    uploaded_file = st.file_uploader("Upload Excel Financial Model", type=["xlsx", "xls", "xlsm"])
with col2:
    observation_notes = st.text_area("Observation Notes", height=100)
    client_requirements = st.text_area("Client-Specific Requirements", height=100)

def smart_excel_extractor(file):
    """
    Searches for CRE keywords in the Excel file. 
    When found, it extracts the table downwards until it hits 3 consecutive blank rows.
    """
    xls = pd.ExcelFile(file)
    extracted_text = []
    
    # Trigger keywords to look for in the model
    target_keywords = [
        't-12', 'trailing 12', 't12', 'rent roll', 'unit mix', 
        'proforma', 'pro forma', 'cash flow', 'return summary', 
        'metrics', 'sources and uses'
    ]
    
    for sheet_name in xls.sheet_names:
        # Read without headers so we can search every single cell
        df = pd.read_excel(xls, sheet_name=sheet_name, header=None)
        
        skip_until = -1
        
        for i, row in df.iterrows():
            if i < skip_until:
                continue
                
            # Convert the row to a single lowercase string to check for keywords
            row_str = ' '.join(row.dropna().astype(str).str.lower())
            
            if any(keyword in row_str for keyword in target_keywords):
                start_idx = i
                end_idx = i
                blank_count = 0
                
                # Scan downwards to find the end of the table
                for j in range(i + 1, len(df)):
                    if df.iloc[j].dropna().empty:
                        blank_count += 1
                    else:
                        blank_count = 0 
                        end_idx = j
                        
                    # Stop if we hit 3 consecutive blank rows
                    if blank_count >= 3: 
                        break
                
                # Extract the chunk and drop entirely empty columns for cleaner reading
                table_chunk = df.iloc[start_idx:end_idx + 1].dropna(how='all', axis=1)
                
                if not table_chunk.empty:
                    extracted_text.append(f"\n--- Found Data in Sheet: '{sheet_name}' ---")
                    # Convert to string. We limit rows per table to 200 just in case it catches a massive raw data dump
                    extracted_text.append(table_chunk.head(200).to_string(index=False, header=False))
                
                skip_until = end_idx + 1 

    # Limit total output length so we don't exceed the LLM's token limit
    final_text = "\n".join(extracted_text)
    return final_text[:50000] 

def create_docx(text):
    """Creates a well-formatted Word document from Markdown text."""
    doc = Document()
    doc.add_heading('Investment Committee Memorandum', 0)
    
    for line in text.split('\n'):
        if line.startswith('### '):
            doc.add_heading(line.replace('### ', '').strip(), level=3)
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', '').strip(), level=2)
        elif line.startswith('# '):
            doc.add_heading(line.replace('# ', '').strip(), level=1)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:].strip(), style='List Bullet')
        elif line.strip():
            doc.add_paragraph(line.strip())
            
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Processing ---
if st.button("Generate IC Memo", type="primary"):
    if not api_key or not uploaded_file:
        st.error("Please provide an API Key and upload an Excel model.")
    else:
        with st.spinner("Running smart extraction and drafting memo..."):
            try:
                # 1. Smart Extraction
                excel_data = smart_excel_extractor(uploaded_file)
                
                if not excel_data.strip():
                    st.warning("Could not automatically find standard tables (T-12, Rent Roll, etc.). The LLM will do its best with limited data.")
                
                # 2. Call Claude API
                client = anthropic.Anthropic(api_key=api_key)
                
                system_prompt = "You are an expert Commercial Real Estate Underwriter drafting an IC Memo."
                
                user_message = f"""
                Draft a professional IC Memo. 
                
                Structure:
                1. Executive Summary
                2. Deal Summary & Headline Metrics
                3. Going-In Cap Rate Analysis
                4. Key Investment Highlights
                5. Unit Mix & Underwriting Assumptions
                
                Client Requirements: {client_requirements}
                Observation Notes: {observation_notes}
                
                Extracted Excel Data:
                {excel_data}
                
                Format using Markdown headers (#, ##) and bullet points.
                """
                
                response = client.messages.create(
                    model="claude-3-5-sonnet-20241022",
                    max_tokens=4000,
                    system=system_prompt,
                    messages=[{"role": "user", "content": user_message}]
                )
                
                memo_content = response.content[0].text
                
                st.divider()
                st.subheader("Generated IC Memo")
                st.markdown(memo_content)
                
                doc_buffer = create_docx(memo_content)
                st.download_button("📄 Download as Word Document", data=doc_buffer, file_name="IC_Memo.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                
            except Exception as e:
                st.error(f"Error: {e}")
